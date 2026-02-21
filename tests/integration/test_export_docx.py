from __future__ import annotations

import base64
import hashlib
import io
import subprocess
import zipfile
from pathlib import Path

from docx import Document

from exam_helper.export_docx import export_project_to_docx, render_project_docx_bytes
from exam_helper.models import FigureData, MCChoice, Question, QuestionType
from exam_helper.repository import ProjectRepository


def _document_xml_from_bytes(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def test_export_docx_with_embedded_figure(tmp_path: Path) -> None:
    repo = ProjectRepository(tmp_path)
    repo.init_project("Exam", "Physics")

    # 1x1 transparent PNG
    b64 = (
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
        "/w8AAgMBgU7Y5e0AAAAASUVORK5CYII="
    )
    raw_hash = hashlib.sha256(base64.b64decode(b64.encode("ascii"))).hexdigest()
    fig = FigureData(
        id="fig_1",
        mime_type="image/png",
        data_base64=b64,
        sha256=raw_hash,
        caption="tiny",
    )
    q = Question(
        id="q1",
        title="t",
        prompt_md="p",
        points=5,
        figures=[fig],
        question_type=QuestionType.multiple_choice,
        choices=[
            MCChoice(label="A", content_md="A1", is_correct=True),
            MCChoice(label="B", content_md="B1", is_correct=False),
            MCChoice(label="C", content_md="C1", is_correct=False),
            MCChoice(label="D", content_md="D1", is_correct=False),
            MCChoice(label="E", content_md="E1", is_correct=False),
        ],
    )
    repo.save_question(q)

    out = tmp_path / "exam.docx"
    warnings = export_project_to_docx(tmp_path, out)
    assert warnings == []
    assert out.exists()
    assert out.stat().st_size > 0
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "[5 points] p" in text
    assert "A1" in text
    assert "E1" in text
    assert "Solution:" not in text
    xml = _document_xml_from_bytes(out.read_bytes())
    assert "<w:numPr>" in xml


def test_export_docx_includes_solution_when_enabled(tmp_path: Path) -> None:
    repo = ProjectRepository(tmp_path)
    repo.init_project("Exam", "Physics")
    q = Question(
        id="q1",
        prompt_md="p",
        points=5,
        solution={"worked_solution_md": "Problem (verbatim): p\nLine 1\n\n\nLine 2"},
    )
    repo.save_question(q)
    out = tmp_path / "exam_with_solution.docx"
    export_project_to_docx(tmp_path, out, include_solutions=True)
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Solution:" in text
    assert "Line 1" in text
    assert "Line 2" in text
    assert "Problem (verbatim):" not in text

    solution_paragraph = next(
        p for p in doc.paragraphs if p.text.strip().startswith("Solution")
    )
    assert solution_paragraph.paragraph_format.left_indent is not None
    assert solution_paragraph.runs
    assert solution_paragraph.runs[0].italic is True


def test_export_docx_falls_back_when_pandoc_missing(tmp_path: Path, monkeypatch) -> None:
    repo = ProjectRepository(tmp_path)
    repo.init_project("Fallback Exam", "Physics")
    q = Question(
        id="q1",
        prompt_md="Compute $v$ from \\(a t\\).",
        points=5,
        question_type=QuestionType.multiple_choice,
        choices=[
            MCChoice(label="A", content_md="$1$", is_correct=True),
            MCChoice(label="B", content_md="$2$", is_correct=False),
        ],
        solution={"worked_solution_md": "Problem (verbatim): Prompt\nline2"},
    )
    repo.save_question(q)

    def _raise_missing(*args, **kwargs):
        raise FileNotFoundError("pandoc")

    monkeypatch.setattr("exam_helper.export_docx.subprocess.run", _raise_missing)
    content, warnings = render_project_docx_bytes(tmp_path)
    assert warnings
    assert "Pandoc not available" in warnings[0]

    doc = Document(io.BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "WARNING:" in text
    assert "[5 points] Compute v from a t." in text
    assert "Problem (verbatim):" not in text
    xml = _document_xml_from_bytes(content)
    assert "<w:numPr>" in xml


def test_export_docx_uses_a_paren_markers_for_pandoc_mc_lists(tmp_path: Path, monkeypatch) -> None:
    repo = ProjectRepository(tmp_path)
    repo.init_project("Pandoc MC", "Physics")
    q = Question(
        id="q1",
        prompt_md="Prompt",
        points=5,
        question_type=QuestionType.multiple_choice,
        choices=[
            MCChoice(label="A", content_md="opt A", is_correct=True),
            MCChoice(label="B", content_md="opt B", is_correct=False),
        ],
        solution={"worked_solution_md": "Line 1"},
    )
    repo.save_question(q)

    def _fake_pandoc(cmd, check, capture_output, text, cwd):
        md_path = Path(cmd[1])
        markdown = md_path.read_text(encoding="utf-8")
        assert "A) opt A" in markdown
        assert "B) opt B" in markdown
        out_path = Path(cmd[cmd.index("-o") + 1])
        d = Document()
        d.add_paragraph("[5 points] Prompt")
        d.add_paragraph("opt A")
        d.add_paragraph("opt B")
        d.add_paragraph("Solution:")
        d.add_paragraph("Line 1")
        d.save(out_path)
        return subprocess.CompletedProcess(cmd, 0, "", "")

    monkeypatch.setattr("exam_helper.export_docx.subprocess.run", _fake_pandoc)
    content, warnings = render_project_docx_bytes(tmp_path, include_solutions=True)
    assert warnings == []

    doc = Document(io.BytesIO(content))
    assert any(p.text == "opt A" for p in doc.paragraphs)
