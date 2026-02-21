from __future__ import annotations

import argparse
from pathlib import Path

import uvicorn

from exam_helper.app import create_app
from exam_helper.config import resolve_openai_api_key
from exam_helper.repository import ProjectRepository
from exam_helper.validation import validate_project


def cmd_init(args: argparse.Namespace) -> int:
    root = Path(args.path)
    repo = ProjectRepository(root)
    repo.init_project(name=args.name, course=args.course)
    return 0


def cmd_validate(args: argparse.Namespace) -> int:
    path = Path(args.path)
    if not path.exists():
        raise SystemExit(f"Path does not exist: {path}")
    if path.is_file():
        path = path.parent
    repo = ProjectRepository(path)
    errors = validate_project(repo)
    if errors:
        for err in errors:
            print(f"ERROR: {err}")
        return 1
    print(f"Validation succeeded: {path}")
    return 0


def cmd_export_docx(args: argparse.Namespace) -> int:
    from docx import Document

    repo = ProjectRepository(Path(args.path))
    project = repo.load_project()
    questions = repo.list_questions()
    doc = Document()
    doc.add_heading(project.name, level=1)
    doc.add_paragraph(project.course)
    for i, q in enumerate(questions, start=1):
        doc.add_heading(f"Question {i}: {q.title}", level=2)
        doc.add_paragraph(q.prompt_md)
    doc.save(args.output)
    print(f"Wrote {args.output}")
    return 0


def cmd_serve(args: argparse.Namespace) -> int:
    key = resolve_openai_api_key(args.openai_key)
    if key:
        print("OpenAI key loaded.")
    else:
        print("OpenAI key not configured. AI features will be unavailable.")
    app = create_app(project_root=Path(args.project), openai_key=key)
    uvicorn.run(app, host=args.host, port=args.port)
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="exam-helper")
    sub = parser.add_subparsers(dest="command", required=True)

    p_init = sub.add_parser("init", help="Initialize a new exam project.")
    p_init.add_argument("path")
    p_init.add_argument("--name", default="Example Exam Project")
    p_init.add_argument("--course", default="Calculus-based Intro Physics")
    p_init.set_defaults(func=cmd_init)

    p_serve = sub.add_parser("serve", help="Serve the local web app.")
    p_serve.add_argument("--host", default="127.0.0.1")
    p_serve.add_argument("--port", type=int, default=8000)
    p_serve.add_argument("--project", default=".")
    p_serve.add_argument("--openai-key", default=None)
    p_serve.set_defaults(func=cmd_serve)

    p_validate = sub.add_parser("validate", help="Validate question files.")
    p_validate.add_argument("path")
    p_validate.set_defaults(func=cmd_validate)

    p_export = sub.add_parser("export", help="Export artifacts.")
    export_sub = p_export.add_subparsers(dest="export_command", required=True)
    p_docx = export_sub.add_parser("docx", help="Export DOCX.")
    p_docx.add_argument("path")
    p_docx.add_argument("--output", required=True)
    p_docx.set_defaults(func=cmd_export_docx)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
