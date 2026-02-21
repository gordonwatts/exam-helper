from __future__ import annotations

import base64
import io
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from exam_helper.models import Question, QuestionType
from exam_helper.repository import ProjectRepository

_PANDOC_MISSING_WARNING = (
    "Pandoc not available; DOCX was exported with plain-text math fallback."
)
_PANDOC_FAILED_WARNING = (
    "Pandoc conversion failed; DOCX was exported with plain-text math fallback."
)


def _normalize_for_docx(text: str) -> str:
    # Keep readable fallback by stripping latex delimiters for Word plain text runs.
    return (
        text.replace("$$", "")
        .replace("$", "")
        .replace("\\(", "")
        .replace("\\)", "")
        .replace("\\[", "")
        .replace("\\]", "")
    )


def _compact_lines(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines()]
    compact: list[str] = []
    prev_blank = False
    for line in lines:
        blank = line == ""
        if blank and prev_blank:
            continue
        compact.append(line)
        prev_blank = blank
    return compact


def _strip_problem_verbatim_lines(text: str) -> str:
    kept: list[str] = []
    for line in text.splitlines():
        if re.match(r"^\s*problem\s*\(verbatim\)\s*:", line, flags=re.IGNORECASE):
            continue
        kept.append(line)
    return "\n".join(kept)


def _apply_numbering_to_paragraph(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for old in list(p_pr.findall(qn("w:numPr"))):
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)


def _apply_solution_paragraph_style(paragraph) -> None:
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def _next_numbering_id(numbering, attr_name: str, tag_name: str) -> int:
    ids = [
        int(raw)
        for el in numbering.findall(qn(tag_name))
        if (raw := el.get(qn(attr_name))) is not None
    ]
    return (max(ids) + 1) if ids else 1


def _create_abstract_numbering(numbering, num_fmt: str) -> int:
    abstract_id = _next_numbering_id(numbering, "w:abstractNumId", "w:abstractNum")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), num_fmt)
    level.append(num_format)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)

    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def _create_numbering_instance(numbering, abstract_id: int) -> int:
    num_id = _next_numbering_id(numbering, "w:numId", "w:num")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)

    numbering.append(num)
    return num_id


def _add_numbered_paragraph(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph(text)
    _apply_numbering_to_paragraph(paragraph, num_id)


def _normalize_math_delimiters(text: str) -> str:
    text = re.sub(r"\\\((.*?)\\\)", r"$\1$", text, flags=re.DOTALL)
    text = re.sub(r"\\\[(.*?)\\\]", r"$$\1$$", text, flags=re.DOTALL)
    return text


def _build_project_markdown(
    project_name: str, course: str, questions: list[Question], include_solutions: bool, image_dir: Path
) -> tuple[str, list[str]]:
    warnings: list[str] = []
    lines: list[str] = [f"# {project_name or 'Exam'}", "", course, ""]

    for i, q in enumerate(questions, start=1):
        prompt = _normalize_math_delimiters(q.prompt_md.strip() or (q.title.strip() if q.title else ""))
        lines.append(f"{i}. [{q.points} points] {prompt}")

        for figure in q.figures:
            try:
                raw = base64.b64decode(figure.data_base64.encode("ascii"))
                ext = figure.mime_type.split("/")[-1] if "/" in figure.mime_type else "bin"
                img_name = f"{q.id}_{figure.id}.{ext}"
                img_path = image_dir / img_name
                img_path.write_bytes(raw)
                caption = figure.caption or "Figure"
                lines.append(f"   ![{caption}]({img_name})")
            except Exception:
                lines.append("   [Figure could not be rendered in DOCX]")
                warnings.append(f"Figure '{figure.id}' in question '{q.id}' could not be exported.")

        if q.question_type == QuestionType.multiple_choice and q.choices:
            for choice in sorted(q.choices, key=lambda c: c.label):
                choice_text = _normalize_math_delimiters(choice.content_md)
                lines.append(f"   {choice.label}) {choice_text}")

        if include_solutions:
            lines.append("")
            lines.append("   *Solution:*")
            cleaned_solution = _strip_problem_verbatim_lines(q.solution.worked_solution_md)
            for line in _compact_lines(_normalize_math_delimiters(cleaned_solution)):
                if line:
                    lines.append(f"   {line}")
            if q.solution.rubric:
                lines.append("   *Rubric:*")
                for item in q.solution.rubric:
                    lines.append(f"   - {_normalize_math_delimiters(item)}")
        lines.append("")
    return "\n".join(lines).strip() + "\n", warnings


def _render_docx_with_pandoc(
    project_name: str, course: str, questions: list[Question], include_solutions: bool
) -> tuple[bytes, list[str]]:
    with tempfile.TemporaryDirectory(prefix="exam-helper-export-") as tmp:
        tmp_path = Path(tmp)
        markdown, warnings = _build_project_markdown(
            project_name=project_name,
            course=course,
            questions=questions,
            include_solutions=include_solutions,
            image_dir=tmp_path,
        )
        md_path = tmp_path / "export.md"
        out_path = tmp_path / "export.docx"
        md_path.write_text(markdown, encoding="utf-8")

        subprocess.run(
            ["pandoc", str(md_path), "-o", str(out_path), "--from=markdown+tex_math_dollars+fancy_lists"],
            check=True,
            capture_output=True,
            text=True,
            cwd=tmp,
        )
        return _postprocess_pandoc_docx(out_path.read_bytes(), include_solutions=include_solutions), warnings


def _postprocess_pandoc_docx(content: bytes, include_solutions: bool) -> bytes:
    doc = Document(io.BytesIO(content))
    in_solution_block = False

    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if not stripped:
            if in_solution_block:
                _apply_solution_paragraph_style(paragraph)
            continue

        if re.match(r"^\[[0-9]+\s+points\]\s+", stripped):
            in_solution_block = False

        if include_solutions and stripped.startswith("Solution:"):
            in_solution_block = True
            _apply_solution_paragraph_style(paragraph)
            continue

        if include_solutions and in_solution_block:
            if re.match(r"^\[[0-9]+\s+points\]\s+", stripped):
                in_solution_block = False
            else:
                _apply_solution_paragraph_style(paragraph)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _add_question(
    doc: Document,
    q: Question,
    include_solutions: bool,
    question_num_id: int,
    choice_abstract_id: int,
) -> None:
    question_text = _normalize_for_docx(q.prompt_md.strip() or (q.title.strip() if q.title else ""))
    _add_numbered_paragraph(doc, f"[{q.points} points] {question_text}", question_num_id)

    for figure in q.figures:
        raw = base64.b64decode(figure.data_base64.encode("ascii"))
        stream = io.BytesIO(raw)
        try:
            doc.add_picture(stream)
            if figure.caption:
                doc.add_paragraph(f"Figure: {figure.caption}")
        except Exception:
            doc.add_paragraph("[Figure could not be rendered in DOCX]")

    if q.question_type == QuestionType.multiple_choice and q.choices:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        choice_num_id = _create_numbering_instance(numbering, choice_abstract_id)
        sorted_choices = sorted(q.choices, key=lambda c: c.label)
        for choice in sorted_choices:
            _add_numbered_paragraph(doc, _normalize_for_docx(choice.content_md), choice_num_id)

    if include_solutions:
        p = doc.add_paragraph("Solution:")
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.25)
        cleaned_solution = _strip_problem_verbatim_lines(q.solution.worked_solution_md)
        for line in _compact_lines(_normalize_for_docx(cleaned_solution)):
            if not line:
                continue
            sol = doc.add_paragraph(line)
            _apply_solution_paragraph_style(sol)
        if q.solution.rubric:
            rubric_head = doc.add_paragraph("Rubric")
            _apply_solution_paragraph_style(rubric_head)
            for item in q.solution.rubric:
                rub = doc.add_paragraph(f"- {_normalize_for_docx(item)}")
                _apply_solution_paragraph_style(rub)


def _render_docx_with_python_docx(
    project_name: str, course: str, questions: list[Question], include_solutions: bool
) -> bytes:
    doc = Document()
    doc.add_heading(project_name or "Exam", level=1)
    doc.add_paragraph(course)

    numbering = doc.part.numbering_part.numbering_definitions._numbering
    question_abstract_id = _create_abstract_numbering(numbering, "decimal")
    choice_abstract_id = _create_abstract_numbering(numbering, "upperLetter")
    question_num_id = _create_numbering_instance(numbering, question_abstract_id)

    for q in questions:
        _add_question(
            doc,
            q,
            include_solutions=include_solutions,
            question_num_id=question_num_id,
            choice_abstract_id=choice_abstract_id,
        )

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def render_project_docx_bytes(
    project_root: Path, include_solutions: bool = False
) -> tuple[bytes, list[str]]:
    repo = ProjectRepository(project_root)
    project = repo.load_project()
    questions = repo.list_questions()

    warnings: list[str] = []
    try:
        return _render_docx_with_pandoc(
            project_name=project.name,
            course=project.course,
            questions=questions,
            include_solutions=include_solutions,
        )
    except FileNotFoundError:
        warnings.append(_PANDOC_MISSING_WARNING)
    except subprocess.CalledProcessError:
        warnings.append(_PANDOC_FAILED_WARNING)

    bytes_out = _render_docx_with_python_docx(
        project_name=project.name,
        course=project.course,
        questions=questions,
        include_solutions=include_solutions,
    )
    # Add warning visibly into exported document content.
    if warnings:
        doc = Document(io.BytesIO(bytes_out))
        for warning in reversed(warnings):
            doc.paragraphs[0].insert_paragraph_before(f"WARNING: {warning}")
        stream = io.BytesIO()
        doc.save(stream)
        bytes_out = stream.getvalue()
    return bytes_out, warnings


def export_project_to_docx(
    project_root: Path, output_path: Path, include_solutions: bool = False
) -> list[str]:
    bytes_out, warnings = render_project_docx_bytes(
        project_root=project_root, include_solutions=include_solutions
    )
    output_path.write_bytes(bytes_out)
    return warnings
